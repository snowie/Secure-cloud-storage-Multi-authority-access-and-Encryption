import json
import os
import numpy as np
import streamlit as st
from streamlit import session_state
import pdfplumber
import docx
import pyaes
import random
import pandas as pd
import base64
import hashlib
from pymongo import MongoClient
import smtplib
import string
import re
import datetime

from dotenv import load_dotenv
load_dotenv()

session_state = st.session_state
if "user_index" not in st.session_state:
    st.session_state["user_index"] = 0


@st.cache_resource
def init_connection(): 
    return MongoClient(
        os.getenv("MONGODB_URI")
    )
    
client = init_connection()
db = client["file_data"]
users_collection = db["users"]
print("Connected to MongoDB!")
print(users_collection)

def user_exists(email):
    try:
        user = users_collection.find_one({"email": email})
        if user is not None:
            return True
        return False
    except Exception as e:
        st.error(f"Error checking user: {e}")
        return False


def send_verification_code(email, code):
    SENDER_MAIL_ID = os.getenv("SENDER_MAIL_ID")
    APP_PASSWORD = os.getenv("APP_PASSWORD")
    RECEIVER = email
    server = smtplib.SMTP_SSL("smtp.googlemail.com", 465)
    server.login(SENDER_MAIL_ID, APP_PASSWORD)
    message = f"Subject: Your Verification Code\n\nYour verification code is: {code}"
    server.sendmail(SENDER_MAIL_ID, RECEIVER, message)
    server.quit()
    st.success("Email sent successfully!")
    return True


def generate_verification_code(length=6):
    return "".join(random.choices(string.ascii_uppercase + string.digits, k=length))


def signup(json_file_path="students.json"):
    st.title("Student Signup Page")
    with st.form("signup_form"):
        st.write("Fill in the details below to create an account:")
        name = st.text_input("Name:")
        email = st.text_input("Email:")
        age = st.number_input("Age:", min_value=0, max_value=120)
        sex = st.radio("Sex:", ("Male", "Female", "Other"))
        password = st.text_input("Password:", type="password")
        confirm_password = st.text_input("Confirm Password:", type="password")
        if (
            session_state.get("verification_code") is None
            or session_state.get("verification_time") is None
            or datetime.datetime.now() - session_state.get("verification_time")
            > datetime.timedelta(minutes=5)
        ):
            verification_code = generate_verification_code()
            session_state["verification_code"] = verification_code
            session_state["verification_time"] = datetime.datetime.now()
        if st.form_submit_button("Signup"):
            if not name:
                st.error("Name field cannot be empty.")
            elif not email:
                st.error("Email field cannot be empty.")
            elif not re.match(r"^[\w\.-]+@[\w\.-]+$", email):
                st.error("Invalid email format. Please enter a valid email address.")
            elif user_exists(email):
                st.error(
                    "User with this email already exists. Please choose a different email."
                )
            elif not age:
                st.error("Age field cannot be empty.")
            elif not password or len(password) < 6:  # Minimum password length of 6
                st.error("Password must be at least 6 characters long.")
            elif password != confirm_password:
                st.error("Passwords do not match. Please try again.")
            else:
                verification_code = session_state["verification_code"]
                send_verification_code(email, verification_code)
                entered_code = st.text_input(
                    "Enter the verification code sent to your email:"
                )
                if entered_code == verification_code:
                    user = create_account(name, email, age, sex, password)
                    session_state["logged_in"] = True
                    session_state["user_info"] = user
                    st.success("Signup successful. You are now logged in!")
                elif len(entered_code) == 6 and entered_code != verification_code:
                    st.error("Incorrect verification code. Please try again.")


def check_login(username, password):

    try:
        user = users_collection.find_one({"email": username, "password": password})
        if user is not None:
            session_state["logged_in"] = True
            session_state["user_info"] = user
            st.success("Login successful!")
            return user
        return None
    except Exception as e:
        st.error(f"Error checking login: {e}")
        return None


def initialize_database():
    try:
        # Check if the users collection exists
        if "users" not in db.list_collection_names():
            db.create_collection("users")
    except Exception as e:
        print(f"Error initializing database: {e}")


def create_account(name, email, age, sex, password):
    try:

        email = email.lower()
        password = hashlib.sha256(password.encode()).hexdigest()
        user_info = {
            "name": name,
            "email": email,
            "age": age,
            "sex": sex,
            "password": password,
            "files": None,
        }
        result = users_collection.insert_one(user_info)
        user_info["_id"] = result.inserted_id
        st.success("Account created successfully! You can now login.")
        return user_info
    except Exception as e:
        st.error(f"Error creating account: {e}")
        return None


def login():
    st.title("Login Page")
    username = st.text_input("Email:")
    password = st.text_input("Password:", type="password")
    username = username.lower()
    password = hashlib.sha256(password.encode()).hexdigest()

    login_button = st.button("Login")

    if login_button:
        user = check_login(username, password)
        if user is not None:
            session_state["logged_in"] = True
            session_state["user_info"] = user
        else:
            st.error("Invalid credentials. Please try again.")


def render_dashboard(user_info):
    try:
        st.title(f"Welcome to the Dashboard, {user_info['name']}!")
        st.subheader("User Information:")
        st.write(f"Name: {user_info['name']}")
        st.write(f"Sex: {user_info['sex']}")
        st.write(f"Age: {user_info['age']}")

    except Exception as e:
        st.error(f"Error rendering dashboard: {e}")


def generateKey(user_key, admin_auth, token_auth):
    key = hashlib.sha256(
        user_key.encode("utf-8")
        + admin_auth.encode("utf-8")
        + token_auth.encode("utf-8")
    ).digest()[:16]
    return key


def extract_text(file) -> str:
    if isinstance(file, str):
        file_extension = os.path.splitext(file)[1].lower()
    else:
        file_extension = os.path.splitext(file.name)[1].lower()

    if file_extension == ".pdf":
        if isinstance(file, str):
            with pdfplumber.open(file) as pdf:
                text = "\n".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
        else:
            with pdfplumber.open(file) as pdf:
                text = "\n".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
    elif file_extension == ".docx":
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        if isinstance(file, str):
            with open(file, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            with file as f:
                text = f.read()
    return text


def render_dashboard(user_info, json_file_path="data.json"):
    try:
        st.title(f"Welcome to the Dashboard, {user_info['name']}!")
        st.subheader("User Information:")
        st.write(f"Name: {user_info['name']}")
        st.write(f"Sex: {user_info['sex']}")
        st.write(f"Age: {user_info['age']}")
        st.image("image.jpg")
    except Exception as e:
        st.error(f"Error rendering dashboard: {e}")


def get_keys():
    with st.form("credentials"):
        st.write("Enter the credentials to encrypt the file:")
        user_key = st.text_input("Enter the User key:")
        admin_auth = st.text_input("Enter the Admin Authentication key:")
        token_auth = st.text_input("Enter the Token Authentication key:")
        if st.form_submit_button("Encrypt and Upload"):
            return user_key, admin_auth, token_auth
    return None, None, None


def main(
    json_file_path="data.json",
):
    
    users_collection = db["users"]
    st.sidebar.title("Secure Multi-Factor Authentication")
    page = st.sidebar.radio(
        "Go to",
        (
            "Signup/Login",
            "Dashboard",
            "File Upload",
            "File Download",
        ),
        key="Secure Multi-Factor Authentication",
    )

    if page == "Signup/Login":
        st.title("Signup/Login Page")
        login_or_signup = st.radio(
            "Select an option", ("Login", "Signup"), key="login_signup"
        )
        if login_or_signup == "Login":
            login()
        else:
            signup()

    elif page == "Dashboard":
        if session_state.get("logged_in"):
            render_dashboard(session_state["user_info"])
        else:
            st.warning("Please login/signup to view the dashboard.")

    elif page == "File Upload":
        if session_state.get("logged_in"):
            st.title("File Upload")
            uploaded_file = st.file_uploader(
                "Choose a file", type=["txt", "pdf", "png", "jpg", "jpeg"]
            )
            if uploaded_file is not None:
                file_details = {
                    "filename": uploaded_file.name,
                    "filetype": uploaded_file.type,
                    "filesize": uploaded_file.size,
                }
                st.write("Name: %s" % uploaded_file.name)
                st.write("Type: %s" % uploaded_file.type)
                st.write("Size: %s" % uploaded_file.size)
                st.write("Enter the credentials to encrypt the file:")
                with st.form("credentials"):
                    st.write("Enter the credentials to encrypt the file:")
                    user_key = st.text_input("Enter the User key:")
                    admin_auth = st.text_input("Enter the Admin Authentication key:")
                    token_auth = st.text_input("Enter the Token Authentication key:")
                    if st.form_submit_button("Encrypt and Upload"):
                        if users_collection.find_one(
                            {"email": session_state["user_info"]["email"]}
                        ):
                            user_info = users_collection.find_one(
                                {"email": session_state["user_info"]["email"]}
                            )
                            if user_info["files"] is None:
                                user_info["files"] = []
                            key = generateKey(user_key, admin_auth, token_auth)
                            aes = pyaes.AESModeOfOperationCTR(key)
                            file = base64.b64encode(uploaded_file.read()).decode(
                                "utf-8"
                            )
                            cipher_text = aes.encrypt(file)
                            cipher_text = base64.b64encode(cipher_text).decode("utf-8")
                            current_time = str(np.datetime64("now"))
                            file_name = uploaded_file.name
                            for file in user_info["files"]:
                                if file["file"] == file_name:
                                    file_name = (
                                        file_name.split(".")[0]
                                        + "_1."
                                        + file_name.split(".")[1]
                                    )
                            user_info["files"].append(
                                {
                                    "file": uploaded_file.name,
                                    "data": cipher_text,
                                    "time": current_time,
                                    "sanitized": False,
                                }
                            )
                            users_collection.update_one(
                                {"email": session_state["user_info"]["email"]},
                                {"$set": {"files": user_info["files"]}},
                            )
                            st.success("File uploaded successfully!")
        else:
            st.warning("Please login/signup to access this page.")

    elif page == "File Download":
        if session_state.get("logged_in"):
            st.title("File Download")
            # i = 1
            # user_info = session_state["user_info"]
            # if len(session_state["user_info"]["files"]) == 0:
            #     st.warning("No files uploaded yet.")
            #     return
            # for file in session_state["user_info"]["files"]:
            #     files = []
            #     file_data = {}
            #     file_data["S.No"] = i
            #     file_data["File Name"] = file["file"]
            #     file_data["Upload Time"] = file["time"]
            #     files.append(file_data)
            #     i += 1
            #     st.table(files)
            #     try:
            #         with st.form("credentials2"):
            #             st.write("Enter the credentials to decrypt the file:")
            #             user_key = st.text_input("Enter the User key:")
            #             admin_auth = st.text_input(
            #                 "Enter the Admin Authentication key:"
            #             )
            #             token_auth = st.text_input(
            #                 "Enter the Token Authentication key:"
            #             )
            #             if st.form_submit_button("Decrypt and Download"):
            #                 key = generateKey(user_key, admin_auth, token_auth)
            #                 aes = pyaes.AESModeOfOperationCTR(key)
            #                 data = base64.b64decode(file["data"])
            #                 decrypted_text = aes.decrypt(data).decode("utf-8")
            #                 data = base64.b64decode(decrypted_text)
            #                 with open(file["file"], "wb") as f:
            #                     f.write(data)
            #                     st.success("File downloaded successfully!")
            #     except Exception as e:
            #         st.error(f"Wrong credentials")
            
            users_collection = users_collection.find_one(
                {"email": session_state["user_info"]["email"]}
            )
            if users_collection["files"] is None:
                st.warning("No files uploaded yet.")
                return
            i = 1
            files = []
            for file in users_collection["files"]:
                file_data = {}
                file_data["S.No"] = i
                file_data["File Name"] = file["file"]
                file_data["Upload Time"] = file["time"]
                files.append(file_data)
                i += 1
                st.table(files)
                try:
                    with st.form("credentials2"):
                        st.write("Enter the credentials to decrypt the file:")
                        user_key = st.text_input("Enter the User key:")
                        admin_auth = st.text_input("Enter the Admin Authentication key:")
                        token_auth = st.text_input("Enter the Token Authentication key:")
                        if st.form_submit_button("Decrypt and Download"):
                            key = generateKey(user_key, admin_auth, token_auth)
                            aes = pyaes.AESModeOfOperationCTR(key)
                            file_name = st.text_input("Enter the file name to download:")
                            file = users_collection["files"]
                            for f in file:
                                if f["file"] == file_name:
                                    data = base64.b64decode(f["data"])
                                    decrypted_text = aes.decrypt(data).decode("utf-8")
                                    data = base64.b64decode(decrypted_text)
                                    with open(file_name, "wb") as f:
                                        f.write(data)
                                        st.success("File downloaded successfully!")
                except Exception as e:
                    st.error(f"Wrong credentials")
                    
                            
        else:
            st.warning("Please login/signup to access this page.")


if __name__ == "__main__":

    initialize_database()
    main()
